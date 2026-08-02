#!/usr/bin/env python3
"""
Export the funding record as a Word document for application packages.

Same source as the website (data/projects.json), so the numbers cannot drift
apart. The summary block is computed, never typed — which is exactly the class
of error this replaces.

Usage:
    python3 tools/export_projects_docx.py [--lang de|en] [-o out.docx]

Requires python-docx:  pip install python-docx
"""

import argparse
import datetime
import json
import pathlib
import sys
from collections import OrderedDict

ROOT = pathlib.Path(__file__).resolve().parent.parent
DATA = ROOT / "data" / "projects.json"

T = {
    "de": {
        "title": "Verzeichnis der Drittmittelprojekte",
        "subtitle": "Prof. Dr.-Ing. Jan-Niklas Voigt-Antons",
        "summary": "Zusammenfassung",
        "lead": ("{v} Mio. € in {n} Vorhaben aus {lines} Förderlinien, kontinuierlich seit "
                 "{since} eingeworben, mit laufenden Bewilligungen bis {until}."),
        "h_role": "Rolle", "h_count": "Vorhaben", "h_volume": "Volumen", "sum": "Summe",
        "roles": {"sole": "Alleiniger Antragsteller",
                  "coordinator": "Verbundkoordinator",
                  "subproject": "Teilprojektleiter (eigenes Teilvorhaben beantragt und geführt)",
                  "co": "Mitantragsteller"},
        "led": ("Eigenständig beantragt und geleitet – als alleiniger Antragsteller, "
                "Verbundkoordinator oder Teilprojektleiter – sind {n} der {total} Vorhaben "
                "mit {v} Mio. €. In allen Teilvorhaben habe ich den Antrag selbst verfasst, "
                "das Budget eigenverantwortlich bewirtschaftet und das daraus finanzierte "
                "Personal geführt."),
        "coord": ("Verbundkoordination – die Federführung eines Konsortiums einschließlich "
                  "Antragsaufbau, Konsortialführung, Berichtswesen und Mittelabruf – habe ich "
                  "in {n} Vorhaben übernommen: {list}."),
        "overview": "Projektübersicht",
        "cols": ["Laufzeit", "Vorhaben", "Förderer", "Rolle", "Summe"],
        "others": ("Weitere Projektbeteiligung ohne eigenes Fördervolumen: {list}."),
        "asof": "Stand: {d}",
    },
    "en": {
        "title": "Record of third-party funded projects",
        "subtitle": "Prof. Dr.-Ing. Jan-Niklas Voigt-Antons",
        "summary": "Summary",
        "lead": ("€{v} million across {n} projects from {lines} funding lines, raised "
                 "continuously since {since}, with awards running through {until}."),
        "h_role": "Role", "h_count": "Projects", "h_volume": "Volume", "sum": "Total",
        "roles": {"sole": "Sole applicant",
                  "coordinator": "Consortium coordinator",
                  "subproject": "Subproject lead (own subproject applied for and managed)",
                  "co": "Co-applicant"},
        "led": ("Applied for and led independently – as sole applicant, consortium "
                "coordinator or subproject lead – are {n} of {total} projects, €{v} million. "
                "In every subproject the proposal was written personally, the budget managed "
                "independently and the funded staff led."),
        "coord": ("Consortium coordination – leading a consortium end to end, including "
                  "proposal design, consortium management, reporting and fund calls – was "
                  "taken on in {n} projects: {list}."),
        "overview": "Project overview",
        "cols": ["Period", "Project", "Funder", "Role", "Volume"],
        "others": ("Further project involvement without own funding volume: {list}."),
        "asof": "As of {d}",
    },
}


def period(p):
    return str(p["from"]) if p["from"] == p["to"] else "%d–%d" % (p["from"], p["to"])


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--lang", choices=("de", "en"), default="de")
    ap.add_argument("-o", "--out")
    args = ap.parse_args()

    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        sys.exit("python-docx is required:  pip install python-docx")

    data = json.loads(DATA.read_text(encoding="utf-8"))
    t = T[args.lang]
    meta = data["meta"]
    projects = sorted(data["projects"], key=lambda p: (-p["from"], -p["to"], p["name"]))

    by = OrderedDict((k, [0, 0]) for k in meta["roles"])
    for p in projects:
        by[p["role"]][0] += 1
        by[p["role"]][1] += p["volume"]

    n_all = len(projects)
    v_all = sum(p["volume"] for p in projects)
    led = [p for p in projects if p["role"] in meta["self_led_roles"]]
    v_led = sum(p["volume"] for p in led)
    coord = [p["short"] for p in projects if p["role"] == "coordinator"]
    lines = meta.get("funding_lines", len({p["funder"] for p in projects}))

    def mio(v):
        s = "%.3f" % (v / 1000.0)
        return s.replace(".", ",") if args.lang == "de" else s

    def k(v):
        if args.lang == "de":
            return format(v, ",").replace(",", ".") + " T€"
        return "€" + format(v, ",") + "k"

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    doc.add_heading(t["title"], level=1)
    sub = doc.add_paragraph(t["subtitle"])
    sub.runs[0].bold = True

    doc.add_heading(t["summary"], level=2)
    doc.add_paragraph(t["lead"].format(
        v=mio(v_all), n=n_all, lines=lines,
        since=min(p["from"] for p in projects), until=max(p["to"] for p in projects)))

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Light Grid Accent 1"
    for i, h in enumerate((t["h_role"], t["h_count"], t["h_volume"])):
        tbl.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for role, (n, v) in by.items():
        if not n:
            continue
        c = tbl.add_row().cells
        c[0].text = t["roles"][role]
        c[1].text = str(n)
        c[2].text = k(v)
    c = tbl.add_row().cells
    for i, val in enumerate((t["sum"], str(n_all), k(v_all))):
        c[i].paragraphs[0].add_run(val).bold = True

    doc.add_paragraph()
    doc.add_paragraph(t["led"].format(n=len(led), total=n_all, v=mio(v_led)))
    doc.add_paragraph(t["coord"].format(n=len(coord), list=", ".join(coord)))

    doc.add_heading(t["overview"], level=2)
    tbl2 = doc.add_table(rows=1, cols=5)
    tbl2.style = "Light Grid Accent 1"
    for i, h in enumerate(t["cols"]):
        tbl2.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for p in projects:
        c = tbl2.add_row().cells
        c[0].text = period(p)
        c[1].text = p["name"]
        c[2].text = p["funder"]
        c[3].text = t["roles"][p["role"]].split(" (")[0]
        c[4].text = k(p["volume"])
        c[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for row in tbl2.rows:
        row.cells[0].width = Cm(2.4)
        row.cells[4].width = Cm(2.0)

    others = data.get("without_own_volume", [])
    if others:
        doc.add_paragraph()
        doc.add_paragraph(t["others"].format(list=", ".join(
            "%s (%s, %s)" % (o["name"], period(o), o["role_label"]) for o in others)))

    doc.add_paragraph()
    stamp = doc.add_paragraph(t["asof"].format(d=meta["updated"]))
    stamp.runs[0].italic = True

    out = pathlib.Path(args.out) if args.out else \
        ROOT / ("Drittmittelprojekte_Voigt-Antons_%s.docx" % meta["updated"]
                if args.lang == "de" else
                "Funding-record_Voigt-Antons_%s.docx" % meta["updated"])
    doc.save(out)
    print("wrote %s — %d projects, %s" % (out.name, n_all, k(v_all)))

    unconfirmed = [p["short"] for p in projects if p.get("role_unconfirmed")]
    if unconfirmed:
        print("note: role not yet confirmed for %s" % ", ".join(unconfirmed))


if __name__ == "__main__":
    main()
